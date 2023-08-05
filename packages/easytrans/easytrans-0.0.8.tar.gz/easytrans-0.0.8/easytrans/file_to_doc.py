import os
from docx import Document
from os.path import exists

from easytrans.constant import FAILURE_STATUS, SUCCESS_STATUS, TXT_FORMAT
from easytrans.utils import FileUtil


class TxtToDoc(object):

    @staticmethod
    def single_txt_to_doc(txt_path, doc_name=None, doc_dir=None):
        if not os.path.isfile(txt_path) or not os.path.splitext(txt_path)[1] in TXT_FORMAT:
            return FAILURE_STATUS, "file format error!"

        status, result = TxtToDoc.txt_transfer_doc([txt_path], [doc_name], doc_dir)
        if status == SUCCESS_STATUS and isinstance(result, list) and len(result) > 0:
            return status, result[0]
        else:
            return status, result

    @staticmethod
    def multiple_txt_to_doc_by_file_dir(txt_dir, doc_dir=None, doc_name_list=None):
        try:

            if not isinstance(txt_dir, str) or not os.path.isdir(txt_dir):
                return FAILURE_STATUS, "'{}' is not a directory!".format(txt_dir)

            file_list = FileUtil.get_dir_file_path_list(txt_dir, suffix=TXT_FORMAT)
            status, result = TxtToDoc.multiple_txt_to_doc_by_file_list(file_list, doc_name_list, doc_dir)

        except Exception as e:
            return FAILURE_STATUS, repr(e)

        else:
            return status, result

    @staticmethod
    def multiple_txt_to_doc_by_file_list(txt_path_list=[], doc_name_list=None, doc_dir=None):
        try:
            if not isinstance(txt_path_list, list) or not len(txt_path_list) > 0:
                return FAILURE_STATUS, "file list is none!"

            return TxtToDoc.txt_transfer_doc(txt_path_list, doc_name_list, doc_dir)

        except Exception as e:
            return FAILURE_STATUS, repr(e)

    @staticmethod
    def txt_transfer_doc(txt_path_list=[], doc_name_list=None, doc_dir=None):
        if not isinstance(txt_path_list, list):
            return FAILURE_STATUS, "'{0}' is not a file list!".format(txt_path_list)

        if doc_name_list and (not isinstance(doc_name_list, list) or len(doc_name_list) < len(txt_path_list)):
            return FAILURE_STATUS, "'{0}' is not matched with file list!".format(doc_name_list)

        doc_list = []
        index = 0

        try:
            for fl in txt_path_list:
                if doc_name_list:
                    doc_name = doc_name_list[index]
                    index = index + 1
                else:
                    doc_name = None

                if not os.path.isfile(fl) or not os.path.splitext(fl)[1] in TXT_FORMAT:
                    continue

                if not doc_name:
                    doc_name = "".join(FileUtil.get_file_name_by_path(fl), ".docx")
                else:
                    doc_name = "".join(doc_name, ".docx")

                if not doc_dir:
                    doc_dir = os.path.dirname(fl)

                with open(fl) as txt_obj:
                    title = txt_obj.readline()
                    txt = txt_obj.read()

                if exists(doc_name):
                    doc = Document(doc_name)
                else:
                    doc = Document()

                doc.add_heading(title)
                doc.add_paragraph(txt)
                save_path = os.path.join(doc_dir, doc_name)
                doc.save(save_path)
                doc_list.append(save_path)

        except Exception as e:
            return FAILURE_STATUS, repr(e)

        else:
            return SUCCESS_STATUS, doc_list
